# ============================================================
# Word Document Loader — .docx Files
#
# Word documents (.docx) are ZIP archives containing XML files.
# python-docx unpacks this for us — we get structured access to:
# - Paragraphs (with style names like "Heading 1", "Heading 2", "Normal")
# - Tables (rows and cells)
#
# We use heading styles as natural section boundaries, exactly like
# the HTML loader uses <h1>–<h6> tags. This gives us semantically
# coherent chunks rather than arbitrary fixed-length slices.
#
# WHY NOT JUST EXTRACT ALL TEXT AS ONE BLOB?
# Word documents often have a logical structure: Introduction → Methods
# → Results → Discussion. If we split by headings, a query about
# "results" can retrieve the Results section specifically rather than
# a middle chunk that happens to straddle two sections.
# ============================================================

import logging
from pathlib import Path

from src.ingestion.document import Document

logger = logging.getLogger(__name__)

# Style names that act as section separators (python-docx style naming)
_HEADING_STYLES = {
    "heading 1", "heading 2", "heading 3",
    "heading 4", "heading 5", "heading 6",
    "title",  # Title style also serves as a top-level heading
}


def load_docx(file_path: Path) -> list[Document]:
    """
    Load a Word .docx file and extract text by heading sections.

    Returns one Document per section (heading + its content).
    Tables within a section are converted to pipe-delimited text rows
    and included in the section's Document.

    Args:
        file_path: Path to the .docx file.

    Returns:
        List of Documents, one per section.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError(
            "python-docx is required for Word document loading. "
            "Install with: pip install python-docx"
        )

    documents: list[Document] = []

    try:
        word_doc = DocxDocument(str(file_path))

        # Track current section state
        current_heading = file_path.stem  # Use filename as heading if no h1 found
        current_parts: list[str] = []

        def flush_section(heading: str, parts: list[str]) -> None:
            text = "\n\n".join(p.strip() for p in parts if p.strip())
            if text:
                documents.append(
                    Document(
                        text=text,
                        metadata={
                            "source": file_path.name,
                            "source_path": str(file_path),
                            "type": "docx",
                            "section": heading,
                        },
                    )
                )

        # python-docx exposes paragraphs and tables via iter_block_items
        # We use a simpler approach: iterate document.element children
        # which preserves the reading order of paragraphs and tables.
        for block in _iter_blocks(word_doc):
            if block["type"] == "paragraph":
                style_name = block["style"].lower()
                text = block["text"].strip()

                if style_name in _HEADING_STYLES and text:
                    # Heading encountered — flush current section, start new one
                    flush_section(current_heading, current_parts)
                    current_heading = text
                    current_parts = []
                elif text:
                    current_parts.append(text)

            elif block["type"] == "table":
                # Convert table to a readable text block within current section
                table_text = _table_to_text(block["table"])
                if table_text:
                    current_parts.append(table_text)

        # Flush the final section
        flush_section(current_heading, current_parts)

        if not documents:
            # Fallback: extract all paragraph text if section parsing found nothing
            all_text = "\n\n".join(
                p.text.strip()
                for p in word_doc.paragraphs
                if p.text.strip()
            )
            if all_text:
                documents.append(
                    Document(
                        text=all_text,
                        metadata={
                            "source": file_path.name,
                            "source_path": str(file_path),
                            "type": "docx",
                            "section": "full_document",
                        },
                    )
                )

    except Exception as exc:
        logger.error("Failed to load DOCX file %s: %s", file_path.name, exc)
        raise

    logger.info("Loaded %d sections from %s", len(documents), file_path.name)
    return documents


def _iter_blocks(word_doc):
    """
    Iterate over paragraphs and tables in document order.

    python-docx's word_doc.paragraphs only returns paragraphs,
    and word_doc.tables only returns tables — neither preserves
    the interleaved order. This generator walks the XML to preserve order.

    Yields dicts:
        {"type": "paragraph", "text": str, "style": str}
        {"type": "table", "table": docx.Table}
    """
    from docx.oxml.ns import qn

    body = word_doc.element.body
    for child in body.iterchildren():
        tag = child.tag

        if tag == qn("w:p"):
            # Paragraph element
            para = None
            for p in word_doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            if para is not None:
                yield {
                    "type": "paragraph",
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal",
                }

        elif tag == qn("w:tbl"):
            # Table element
            for table in word_doc.tables:
                if table._element is child:
                    yield {"type": "table", "table": table}
                    break


def _table_to_text(table) -> str:
    """
    Convert a python-docx Table to a pipe-delimited text representation.

    Example output:
        Name | Age | Department
        Alice | 30 | Engineering
        Bob | 25 | Design

    This format is readable as plain text and embeds well.
    """
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)
