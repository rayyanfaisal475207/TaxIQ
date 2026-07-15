import os
import uuid
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Absolute path so downloads survive a server started from any working directory
GENERATED_DIR = str(Path(__file__).resolve().parent.parent.parent / "data" / "generated")

def build_docx(payload: dict) -> tuple[str, int]:
    """
    Builds a DOCX file from the JSON payload.
    Returns (filepath, file_size_bytes).
    """
    os.makedirs(GENERATED_DIR, exist_ok=True)
    filename = f"{uuid.uuid4()}.docx"
    filepath = os.path.join(GENERATED_DIR, filename)

    doc = Document()
    
    # Title
    title = payload.get("title", "TaxIQ Export")
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Description
    desc = payload.get("description", "")
    if desc:
        doc.add_paragraph(desc)
        
    # Sections
    for sec in payload.get("sections", []):
        stype = sec.get("type")
        
        if stype == "heading":
            level = sec.get("level", 2)
            doc.add_heading(sec.get("content", ""), level=level)
            
        elif stype == "paragraph":
            doc.add_paragraph(sec.get("content", ""))
            
        elif stype == "table":
            headers = sec.get("headers", [])
            rows = sec.get("rows", [])
            
            if not headers and not rows:
                continue

            num_cols = len(headers) if headers else max((len(r) for r in rows), default=0)
            if num_cols == 0:
                continue
            table = doc.add_table(rows=1 + len(rows), cols=num_cols)
            table.style = 'Table Grid'
            
            # Add headers
            if headers:
                hdr_cells = table.rows[0].cells
                for i, h in enumerate(headers):
                    hdr_cells[i].text = str(h)
                    # Bold headers
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            
            # Add rows
            for row_idx, rdata in enumerate(rows):
                row_cells = table.rows[row_idx + 1].cells
                for col_idx, cell_data in enumerate(rdata):
                    if col_idx < len(row_cells):
                        row_cells[col_idx].text = str(cell_data)

    doc.save(filepath)
    file_size = os.path.getsize(filepath)
    return filepath, file_size
